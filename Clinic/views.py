from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.shortcuts import render
from django.shortcuts import render, redirect, HttpResponse
from django.contrib.auth import authenticate, login, logout
from openpyxl import Workbook

from .models import *
from .forms import *
import openpyxl
import json
from django.http import JsonResponse
from django.forms.models import model_to_dict
from django.db.models import Q
from docx import Document
from docx.shared import Inches

from functools import wraps


def doctors_only(function):
    @wraps(function)
    def wrap(request, *args, **kwargs):

        profile = request.user
        if profile.username and profile.status and profile.status == 'doctor':
            return function(request, *args, **kwargs)
        else:
            return redirect('main')

    return wrap


def admins_only(function):
    @wraps(function)
    def wrap(request, *args, **kwargs):
        profile = request.user
        if profile.username and profile.status and profile.status == 'admin':
            return function(request, *args, **kwargs)
        else:
            return redirect('main')

    return wrap


def log_in(request):
    if request.POST:
        if request.user.username:
            if request.user.status == 'admin':
                return redirect('admin')
            elif request.user.status == 'doctor':
                return redirect('doctor')
            else:
                return HttpResponse(
                    "Siz hali saytdan foydalana olmaysiz! Iltimos O'zingiz ishlayotgan korxonaga yoki sayt Mualliflariga murojaat qiling")
        username = request.POST.get('username')
        password = request.POST.get('password')
        if request.user.username == '':
            user = authenticate(request, username=username, password=password)
            if user:
                login(request=request, user=user)
                if user.status == 'admin':
                    return redirect('admin')
                elif user.status == 'doctor':
                    return redirect('doctor')
                else:
                    return HttpResponse("Sizda hali status yo'q")
            else:
                return HttpResponse("Username yoki password xato")
    return render(request, 'login.html')


def logoutView(request):
    logout(request)
    return redirect('main')


@admins_only
def adminView(request):
    return render(request, 'Admin main page.html', {})


@admins_only
def EtalonView(request):
    if request.user.username and request.user.status == 'admin':
        form = ExcelFileUploudForm()
        context = {
            'status': "Илтимос EXCEL форматдаги файлни юкланг",
            'form': form
        }
        if request.method == "POST":
            form = ExcelFileUploudForm(request.POST, request.FILES)
        if request.FILES and form.is_valid():
            newdoc = ExcelFile()
            newdoc.user = request.user
            newdoc.file = request.FILES.get('myfile')
            newdoc.save()
            path = newdoc.file.path
            wb_obj = openpyxl.load_workbook(path)
            sanoq_class = 1
            if wb_obj:
                Etalon.objects.all().delete()
            for i in wb_obj.worksheets:

                sheet_name = i.title
                if sheet_name[:5] == "Class":
                    clas = Clas.objects.get(name=f"{sanoq_class}-class")
                    sanoq_class += 1
                    sanoq = 1
                for j in i.iter_rows():
                    if sanoq > 2:
                        etalon = Etalon()
                        etalon.clas = clas
                        etalon.save()
                        etalon.x1 = j[1].value
                        etalon.x2 = j[2].value
                        etalon.x3 = j[3].value
                        etalon.x4 = j[4].value
                        etalon.x5 = j[5].value
                        etalon.x6 = j[6].value
                        etalon.x7 = j[7].value
                        etalon.x8 = j[8].value
                        etalon.x9 = j[9].value
                        etalon.x10 = j[10].value
                        etalon.x11 = j[11].value
                        etalon.x12 = j[12].value
                        etalon.x13 = j[13].value
                        etalon.x14 = j[14].value
                        etalon.x15 = j[15].value
                        etalon.x16 = j[16].value
                        etalon.x17 = j[17].value
                        etalon.x18 = j[18].value
                        etalon.x19 = j[19].value
                        etalon.x20 = j[20].value
                        etalon.x21 = j[21].value
                        etalon.x22 = j[22].value
                        etalon.x23 = j[23].value
                        etalon.x24 = j[24].value
                        etalon.x25 = j[25].value
                        etalon.x26 = j[26].value
                        etalon.x27 = j[27].value
                        etalon.x28 = j[28].value
                        etalon.x29 = j[29].value
                        etalon.x30 = j[30].value
                        etalon.x31 = j[31].value
                        etalon.x32 = j[32].value
                        etalon.x33 = j[33].value
                        etalon.x34 = j[34].value
                        etalon.x35 = j[35].value
                        etalon.x36 = j[36].value
                        etalon.x37 = j[37].value
                        etalon.x38 = j[38].value
                        etalon.x39 = j[39].value
                        etalon.x40 = j[40].value
                        etalon.x41 = j[41].value
                        etalon.x42 = j[42].value
                        etalon.x43 = j[43].value
                        etalon.x44 = j[44].value
                        etalon.x45 = j[45].value
                        etalon.x46 = j[46].value
                        etalon.x47 = j[47].value
                        etalon.x48 = j[48].value
                        etalon.x49 = j[49].value
                        etalon.x50 = j[50].value
                        etalon.x51 = j[51].value
                        etalon.x52 = j[52].value
                        etalon.x53 = j[53].value
                        etalon.x54 = j[54].value
                        etalon.x55 = j[55].value
                        etalon.x56 = j[56].value
                        etalon.x57 = j[57].value
                        etalon.x58 = j[58].value
                        etalon.x59 = j[59].value
                        etalon.x60 = j[60].value
                        etalon.x61 = j[61].value
                        etalon.x62 = j[62].value
                        etalon.x63 = j[63].value
                        etalon.x64 = j[64].value
                        etalon.x65 = j[65].value
                        etalon.x66 = j[66].value
                        etalon.x67 = j[67].value
                        etalon.x68 = j[68].value
                        etalon.x69 = j[69].value
                        etalon.x70 = j[70].value
                        etalon.x71 = j[71].value
                        etalon.x72 = j[72].value
                        etalon.x73 = j[73].value
                        etalon.x74 = j[74].value
                        etalon.x75 = j[75].value
                        etalon.x76 = j[76].value
                        etalon.x77 = j[77].value
                        etalon.x78 = j[78].value
                        etalon.x79 = j[79].value
                        etalon.x80 = j[80].value
                        etalon.x81 = j[81].value
                        etalon.x82 = j[82].value
                        etalon.x83 = j[83].value
                        etalon.x84 = j[84].value
                        etalon.x85 = j[85].value
                        etalon.x86 = j[86].value
                        etalon.x87 = j[87].value
                        etalon.x88 = j[88].value
                        etalon.x89 = j[89].value
                        etalon.save()
                    sanoq += 1
            context = {
                'status': 'Fayl muaffaqiyatli yuklandi'

            }
        return render(request, 'index.html', context)
    else:
        if request.user.username:
            logout(request)
            return redirect('login')

        else:
            return redirect('signup')


def main(request):
    creators = Creator.objects.all()
    context = {
        'creators': creators
    }
    return render(request, 'main.html', context)


@doctors_only
def doctorView(request):
    return render(request, 'Doctor.html')


def signUpView(request):
    form = CustomUserCreationForm(request.POST)
    if form.is_valid():
        form.save()
        username = request.POST.get("username")
        password = request.POST.get("password1")
        user = authenticate(request=request, username=username, password=password)
        if user:
            return redirect('login')
    context = {
        'form': form,
    }
    return render(request, 'signup.html', context)


def classRenderView(request):
    data = json.loads(request.body)
    class_name = data['class']

    clas = Clas.objects.get(name=class_name)
    etalons = Etalon.objects.filter(clas=clas)
    data = []
    for i in etalons:
        data.append({
            'x1': i.x1,
            'x2': i.x2,
            'x3': i.x3,
            'x4': i.x4,
            'x5': i.x5,
            'x6': i.x6,
            'x7': i.x7,
            'x8': i.x8,
            'x9': i.x9,
            'x10': i.x10,
            'x11': i.x11,
            'x12': i.x12,
            'x13': i.x13,
            'x14': i.x14,
            'x15': i.x15,
            'x16': i.x16,
            'x17': i.x17,
            'x18': i.x18,
            'x19': i.x19,
            'x20': i.x20,
            'x21': i.x21,
            'x22': i.x22,
            'x23': i.x23,
            'x24': i.x24,
            'x25': i.x25,
            'x26': i.x26,
            'x27': i.x27,
            'x28': i.x28,
            'x29': i.x29,
            'x30': i.x30,
            'x31': i.x31,
            'x32': i.x32,
            'x33': i.x33,
            'x34': i.x34,
            'x35': i.x35,
            'x36': i.x36,
            'x37': i.x37,
            'x38': i.x38,
            'x39': i.x39,
            'x40': i.x40,
            'x41': i.x41,
            'x42': i.x42,
            'x43': i.x43,
            'x44': i.x44,
            'x45': i.x45,
            'x46': i.x46,
            'x47': i.x47,
            'x48': i.x48,
            'x49': i.x49,
            'x50': i.x50,
            'x51': i.x51,
            'x52': i.x52,
            'x53': i.x53,
            'x54': i.x54,
            'x55': i.x55,
            'x56': i.x56,
            'x57': i.x57,
            'x58': i.x58,
            'x59': i.x59,
            'x60': i.x60,
            'x61': i.x61,
            'x62': i.x62,
            'x63': i.x63,
            'x64': i.x64,
            'x65': i.x65,
            'x66': i.x66,
            'x67': i.x67,
            'x68': i.x68,
            'x69': i.x69,
            'x70': i.x70,
            'x71': i.x71,
            'x72': i.x72,
            'x73': i.x73,
            'x74': i.x74,
            'x75': i.x75,
            'x76': i.x76,
            'x77': i.x77,
            'x78': i.x78,
            'x79': i.x79,
            'x80': i.x80,
            'x81': i.x81,
            'x82': i.x82,
            'x83': i.x83,
            'x84': i.x84,
            'x85': i.x85,
            'x86': i.x86,
            'x87': i.x87,
            'x88': i.x88,
            'x89': i.x89,
        })
    return JsonResponse({'data': data})


@admins_only
def simptokompleksView(request):
    if request.user.username and request.user.status == 'admin':
        data = [
            {'L': Simptokompleks.objects.all(),
             'id': 10},
        ]
        form = SimptokompleksForm()
        context = {
            'status': "Iltimos excel formatda malumot yuklang",
            'form': form,
            'simptokompleksdata': data
        }
        if request.method == "POST":
            form = SimptokompleksForm(request.POST, request.FILES)
        if form.is_valid() and request.FILES:
            excelfile = ExcelFile.objects.create(file=request.FILES.get('myfile'))
            excelfile.save()
            wb_obj = openpyxl.load_workbook(excelfile.file.path)
            if wb_obj:
                Simptokompleks.objects.all().delete()
                Lsimptokompleks.objects.all().delete()
            for i in wb_obj.worksheets:
                sheet_name = i.title
                sanoq = 1
                for j in i.iter_rows():
                    if sanoq > 2:
                        etalon = Simptokompleks()
                        etalon.file = request.FILES.get('myfile')
                        etalon.save()
                        etalon.x1 = j[1].value
                        etalon.x2 = j[2].value
                        etalon.x3 = j[3].value
                        etalon.x4 = j[4].value
                        etalon.x5 = j[5].value
                        etalon.x6 = j[6].value
                        etalon.x7 = j[7].value
                        etalon.x8 = j[8].value
                        etalon.x9 = j[9].value
                        etalon.x10 = j[10].value
                        etalon.x11 = j[11].value
                        etalon.x12 = j[12].value
                        etalon.x13 = j[13].value
                        etalon.x14 = j[14].value
                        etalon.x15 = j[15].value
                        etalon.x16 = j[16].value
                        etalon.x17 = j[17].value
                        etalon.x18 = j[18].value
                        etalon.x19 = j[19].value
                        etalon.x20 = j[20].value
                        etalon.x21 = j[21].value
                        etalon.x22 = j[22].value
                        etalon.x23 = j[23].value
                        etalon.x24 = j[24].value
                        etalon.x25 = j[25].value
                        etalon.x26 = j[26].value
                        etalon.x27 = j[27].value
                        etalon.x28 = j[28].value
                        etalon.x29 = j[29].value
                        etalon.x30 = j[30].value
                        etalon.x31 = j[31].value
                        etalon.x32 = j[32].value
                        etalon.x33 = j[33].value
                        etalon.x34 = j[34].value
                        etalon.x35 = j[35].value
                        etalon.x36 = j[36].value
                        etalon.x37 = j[37].value
                        etalon.x38 = j[38].value
                        etalon.x39 = j[39].value
                        etalon.x40 = j[40].value
                        etalon.x41 = j[41].value
                        etalon.x42 = j[42].value
                        etalon.x43 = j[43].value
                        etalon.x44 = j[44].value
                        etalon.x45 = j[45].value
                        etalon.x46 = j[46].value
                        etalon.x47 = j[47].value
                        etalon.x48 = j[48].value
                        etalon.x49 = j[49].value
                        etalon.x50 = j[50].value
                        etalon.x51 = j[51].value
                        etalon.x52 = j[52].value
                        etalon.x53 = j[53].value
                        etalon.x54 = j[54].value
                        etalon.x55 = j[55].value
                        etalon.x56 = j[56].value
                        etalon.x57 = j[57].value
                        etalon.x58 = j[58].value
                        etalon.x59 = j[59].value
                        etalon.x60 = j[60].value
                        etalon.x61 = j[61].value
                        etalon.x62 = j[62].value
                        etalon.x63 = j[63].value
                        etalon.x64 = j[64].value
                        etalon.x65 = j[65].value
                        etalon.x66 = j[66].value
                        etalon.x67 = j[67].value
                        etalon.x68 = j[68].value
                        etalon.x69 = j[69].value
                        etalon.x70 = j[70].value
                        etalon.x71 = j[71].value
                        etalon.x72 = j[72].value
                        etalon.x73 = j[73].value
                        etalon.x74 = j[74].value
                        etalon.x75 = j[75].value
                        etalon.x76 = j[76].value
                        etalon.x77 = j[77].value
                        etalon.x78 = j[78].value
                        etalon.x79 = j[79].value
                        etalon.x80 = j[80].value
                        etalon.x81 = j[81].value
                        etalon.x82 = j[82].value
                        etalon.x83 = j[83].value
                        etalon.x84 = j[84].value
                        etalon.x85 = j[85].value
                        etalon.x86 = j[86].value
                        etalon.x87 = j[87].value
                        etalon.x88 = j[88].value
                        etalon.x89 = j[89].value
                        etalon.save()
                    sanoq += 1

            context = {
                'status': 'Fayl muaffaqiyatli yuklandi',
                'simptokompleksdata': data

            }
        return render(request, 'simptokompleks.html', context)
    else:
        if request.user.username:
            logout(request)
            return redirect('login')

        else:
            return redirect('signup')


def nolbirlashtirish(clas, patient_id):
    etalon = Etalon.objects.filter(clas__name=clas)
    dataetalon = []
    etalondatas = []
    res = []
    for i in etalon:
        res.append(i.id)
        res.append(i.normallash.x1)
        res.append(i.normallash.x2)
        res.append(i.normallash.x3)
        res.append(i.normallash.x4)
        res.append(i.normallash.x5)
        res.append(i.normallash.x6)
        res.append(i.normallash.x7)
        res.append(i.normallash.x8)
        res.append(i.normallash.x9)
        res.append(i.normallash.x10)
        res.append(i.normallash.x11)
        res.append(i.normallash.x12)
        res.append(i.normallash.x13)
        res.append(i.normallash.x14)
        res.append(i.normallash.x15)
        res.append(i.normallash.x16)
        res.append(i.normallash.x17)
        res.append(i.normallash.x18)
        res.append(i.normallash.x19)
        res.append(i.normallash.x20)
        res.append(i.normallash.x21)
        res.append(i.normallash.x22)
        res.append(i.normallash.x23)
        res.append(i.normallash.x24)
        res.append(i.normallash.x25)
        res.append(i.normallash.x26)
        res.append(i.normallash.x27)
        res.append(i.normallash.x28)
        res.append(i.normallash.x29)
        res.append(i.normallash.x30)
        res.append(i.normallash.x31)
        res.append(i.normallash.x32)
        res.append(i.normallash.x33)
        res.append(i.normallash.x34)
        res.append(i.normallash.x35)
        res.append(i.normallash.x36)
        res.append(i.normallash.x37)
        res.append(i.normallash.x38)
        res.append(i.normallash.x39)
        res.append(i.normallash.x40)
        res.append(i.normallash.x41)
        res.append(i.normallash.x42)
        res.append(i.normallash.x43)
        res.append(i.normallash.x44)
        res.append(i.normallash.x45)
        res.append(i.normallash.x46)
        res.append(i.normallash.x47)
        res.append(i.normallash.x48)
        res.append(i.normallash.x49)
        res.append(i.normallash.x50)
        res.append(i.normallash.x51)
        res.append(i.normallash.x52)
        res.append(i.normallash.x53)
        res.append(i.normallash.x54)
        res.append(i.normallash.x55)
        res.append(i.normallash.x56)
        res.append(i.normallash.x57)
        res.append(i.normallash.x58)
        res.append(i.normallash.x59)
        res.append(i.normallash.x60)
        res.append(i.normallash.x61)
        res.append(i.normallash.x62)
        res.append(i.normallash.x63)
        res.append(i.normallash.x64)
        res.append(i.normallash.x65)
        res.append(i.normallash.x66)
        res.append(i.normallash.x67)
        res.append(i.normallash.x68)
        res.append(i.normallash.x69)
        res.append(i.normallash.x70)
        res.append(i.normallash.x71)
        res.append(i.normallash.x72)
        res.append(i.normallash.x73)
        res.append(i.normallash.x74)
        res.append(i.normallash.x75)
        res.append(i.normallash.x76)
        res.append(i.normallash.x77)
        res.append(i.normallash.x78)
        res.append(i.normallash.x79)
        res.append(i.normallash.x80)
        res.append(i.normallash.x81)
        res.append(i.normallash.x82)
        res.append(i.normallash.x83)
        res.append(i.normallash.x84)
        res.append(i.normallash.x85)
        res.append(i.normallash.x86)
        res.append(i.normallash.x87)
        res.append(i.normallash.x88)
        res.append(i.normallash.x89)
        etalondatas.append(res)
        res = []
        dataetalon.append({
            'x1': i.normallash.x1,
            'x2': i.normallash.x2,
            'x3': i.normallash.x3,
            'x4': i.normallash.x4,
            'x5': i.normallash.x5,
            'x6': i.normallash.x6,
            'x7': i.normallash.x7,
            'x8': i.normallash.x8,
            'x9': i.normallash.x9,
            'x10': i.normallash.x10,
            'x11': i.normallash.x11,
            'x12': i.normallash.x12,
            'x13': i.normallash.x13,
            'x14': i.normallash.x14,
            'x15': i.normallash.x15,
            'x16': i.normallash.x16,
            'x17': i.normallash.x17,
            'x18': i.normallash.x18,
            'x19': i.normallash.x19,
            'x20': i.normallash.x20,
            'x21': i.normallash.x21,
            'x22': i.normallash.x22,
            'x23': i.normallash.x23,
            'x24': i.normallash.x24,
            'x25': i.normallash.x25,
            'x26': i.normallash.x26,
            'x27': i.normallash.x27,
            'x28': i.normallash.x28,
            'x29': i.normallash.x29,
            'x30': i.normallash.x30,
            'x31': i.normallash.x31,
            'x32': i.normallash.x32,
            'x33': i.normallash.x33,
            'x34': i.normallash.x34,
            'x35': i.normallash.x35,
            'x36': i.normallash.x36,
            'x37': i.normallash.x37,
            'x38': i.normallash.x38,
            'x39': i.normallash.x39,
            'x40': i.normallash.x40,
            'x41': i.normallash.x41,
            'x42': i.normallash.x42,
            'x43': i.normallash.x43,
            'x44': i.normallash.x44,
            'x45': i.normallash.x45,
            'x46': i.normallash.x46,
            'x47': i.normallash.x47,
            'x48': i.normallash.x48,
            'x49': i.normallash.x49,
            'x50': i.normallash.x50,
            'x51': i.normallash.x51,
            'x52': i.normallash.x52,
            'x53': i.normallash.x53,
            'x54': i.normallash.x54,
            'x55': i.normallash.x55,
            'x56': i.normallash.x56,
            'x57': i.normallash.x57,
            'x58': i.normallash.x58,
            'x59': i.normallash.x59,
            'x60': i.normallash.x60,
            'x61': i.normallash.x61,
            'x62': i.normallash.x62,
            'x63': i.normallash.x63,
            'x64': i.normallash.x64,
            'x65': i.normallash.x65,
            'x66': i.normallash.x66,
            'x67': i.normallash.x67,
            'x68': i.normallash.x68,
            'x69': i.normallash.x69,
            'x70': i.normallash.x70,
            'x71': i.normallash.x71,
            'x72': i.normallash.x72,
            'x73': i.normallash.x73,
            'x74': i.normallash.x74,
            'x75': i.normallash.x75,
            'x76': i.normallash.x76,
            'x77': i.normallash.x77,
            'x78': i.normallash.x78,
            'x79': i.normallash.x79,
            'x80': i.normallash.x80,
            'x81': i.normallash.x81,
            'x82': i.normallash.x82,
            'x83': i.normallash.x83,
            'x84': i.normallash.x84,
            'x85': i.normallash.x85,
            'x86': i.normallash.x86,
            'x87': i.normallash.x87,
            'x88': i.normallash.x88,
            'x89': i.normallash.x89,
        })

    patients = Patient.objects.filter(id=patient_id)
    datapatient = []
    patientdatas = []
    res = []
    for i in patients:
        res.append(f"{i.last_name} {i.first_name}_{i.id}")
        res.append(i.normallash.x1)
        res.append(i.normallash.x2)
        res.append(i.normallash.x3)
        res.append(i.normallash.x4)
        res.append(i.normallash.x5)
        res.append(i.normallash.x6)
        res.append(i.normallash.x7)
        res.append(i.normallash.x8)
        res.append(i.normallash.x9)
        res.append(i.normallash.x10)
        res.append(i.normallash.x11)
        res.append(i.normallash.x12)
        res.append(i.normallash.x13)
        res.append(i.normallash.x14)
        res.append(i.normallash.x15)
        res.append(i.normallash.x16)
        res.append(i.normallash.x17)
        res.append(i.normallash.x18)
        res.append(i.normallash.x19)
        res.append(i.normallash.x20)
        res.append(i.normallash.x21)
        res.append(i.normallash.x22)
        res.append(i.normallash.x23)
        res.append(i.normallash.x24)
        res.append(i.normallash.x25)
        res.append(i.normallash.x26)
        res.append(i.normallash.x27)
        res.append(i.normallash.x28)
        res.append(i.normallash.x29)
        res.append(i.normallash.x30)
        res.append(i.normallash.x31)
        res.append(i.normallash.x32)
        res.append(i.normallash.x33)
        res.append(i.normallash.x34)
        res.append(i.normallash.x35)
        res.append(i.normallash.x36)
        res.append(i.normallash.x37)
        res.append(i.normallash.x38)
        res.append(i.normallash.x39)
        res.append(i.normallash.x40)
        res.append(i.normallash.x41)
        res.append(i.normallash.x42)
        res.append(i.normallash.x43)
        res.append(i.normallash.x44)
        res.append(i.normallash.x45)
        res.append(i.normallash.x46)
        res.append(i.normallash.x47)
        res.append(i.normallash.x48)
        res.append(i.normallash.x49)
        res.append(i.normallash.x50)
        res.append(i.normallash.x51)
        res.append(i.normallash.x52)
        res.append(i.normallash.x53)
        res.append(i.normallash.x54)
        res.append(i.normallash.x55)
        res.append(i.normallash.x56)
        res.append(i.normallash.x57)
        res.append(i.normallash.x58)
        res.append(i.normallash.x59)
        res.append(i.normallash.x60)
        res.append(i.normallash.x61)
        res.append(i.normallash.x62)
        res.append(i.normallash.x63)
        res.append(i.normallash.x64)
        res.append(i.normallash.x65)
        res.append(i.normallash.x66)
        res.append(i.normallash.x67)
        res.append(i.normallash.x68)
        res.append(i.normallash.x69)
        res.append(i.normallash.x70)
        res.append(i.normallash.x71)
        res.append(i.normallash.x72)
        res.append(i.normallash.x73)
        res.append(i.normallash.x74)
        res.append(i.normallash.x75)
        res.append(i.normallash.x76)
        res.append(i.normallash.x77)
        res.append(i.normallash.x78)
        res.append(i.normallash.x79)
        res.append(i.normallash.x80)
        res.append(i.normallash.x81)
        res.append(i.normallash.x82)
        res.append(i.normallash.x83)
        res.append(i.normallash.x84)
        res.append(i.normallash.x85)
        res.append(i.normallash.x86)
        res.append(i.normallash.x87)
        res.append(i.normallash.x88)
        res.append(i.normallash.x89)
        patientdatas.append(res)
        res = []
        datapatient.append({
            'name': i.first_name,
            'x1': i.normallash.x1,
            'x2': i.normallash.x2,
            'x3': i.normallash.x3,
            'x4': i.normallash.x4,
            'x5': i.normallash.x5,
            'x6': i.normallash.x6,
            'x7': i.normallash.x7,
            'x8': i.normallash.x8,
            'x9': i.normallash.x9,
            'x10': i.normallash.x10,
            'x11': i.normallash.x11,
            'x12': i.normallash.x12,
            'x13': i.normallash.x13,
            'x14': i.normallash.x14,
            'x15': i.normallash.x15,
            'x17': i.normallash.x17,
            'x18': i.normallash.x18,
            'x19': i.normallash.x19,
            'x20': i.normallash.x20,
            'x21': i.normallash.x21,
            'x22': i.normallash.x22,
            'x23': i.normallash.x23,
            'x24': i.normallash.x24,
            'x25': i.normallash.x25,
            'x26': i.normallash.x26,
            'x27': i.normallash.x27,
            'x28': i.normallash.x28,
            'x29': i.normallash.x29,
            'x30': i.normallash.x30,
            'x31': i.normallash.x31,
            'x32': i.normallash.x32,
            'x33': i.normallash.x33,
            'x34': i.normallash.x34,
            'x35': i.normallash.x35,
            'x36': i.normallash.x36,
            'x37': i.normallash.x37,
            'x38': i.normallash.x38,
            'x39': i.normallash.x39,
            'x40': i.normallash.x40,
            'x41': i.normallash.x41,
            'x42': i.normallash.x42,
            'x43': i.normallash.x43,
            'x44': i.normallash.x44,
            'x45': i.normallash.x45,
            'x46': i.normallash.x46,
            'x47': i.normallash.x47,
            'x48': i.normallash.x48,
            'x49': i.normallash.x49,
            'x50': i.normallash.x50,
            'x51': i.normallash.x51,
            'x52': i.normallash.x52,
            'x53': i.normallash.x53,
            'x54': i.normallash.x54,
            'x55': i.normallash.x55,
            'x56': i.normallash.x56,
            'x57': i.normallash.x57,
            'x58': i.normallash.x58,
            'x59': i.normallash.x59,
            'x60': i.normallash.x60,
            'x61': i.normallash.x61,
            'x62': i.normallash.x62,
            'x63': i.normallash.x63,
            'x64': i.normallash.x64,
            'x65': i.normallash.x65,
            'x66': i.normallash.x66,
            'x67': i.normallash.x67,
            'x68': i.normallash.x68,
            'x69': i.normallash.x69,
            'x70': i.normallash.x70,
            'x71': i.normallash.x71,
            'x72': i.normallash.x72,
            'x73': i.normallash.x73,
            'x74': i.normallash.x74,
            'x75': i.normallash.x75,
            'x76': i.normallash.x76,
            'x77': i.normallash.x77,
            'x78': i.normallash.x78,
            'x79': i.normallash.x79,
            'x80': i.normallash.x80,
            'x81': i.normallash.x81,
            'x82': i.normallash.x82,
            'x83': i.normallash.x83,
            'x84': i.normallash.x84,
            'x85': i.normallash.x85,
            'x86': i.normallash.x86,
            'x87': i.normallash.x87,
            'x88': i.normallash.x88,
            'x89': i.normallash.x89,
        })
    middledata = []
    adata = []
    for i in range(90):
        middledata.append(0)
        adata.append(0)
    for i in etalondatas:
        for j in range(1, 90):
            middledata[j] += i[j]
    for i in patientdatas:
        for j in range(1, 90):
            middledata[j] += i[j]
    for i in range(1, 90):
        middledata[i] = (middledata[i]) / (len(etalon) + len(patients))
    for i in etalondatas:
        for j in range(1, 90):
            adata[j] += ((middledata[j] - i[j]) ** 2)
    for i in patientdatas:
        for j in range(1, 90):
            adata[j] += (middledata[j] - i[j]) ** 2
    soni = len(etalondatas) + len(patientdatas)
    for i in range(1, 90):
        adata[i] = adata[i] / soni
    datab = []
    res = []
    for i in etalondatas:
        res.append(i[0])
        for j in range(1, 90):
            res.append((middledata[j] - i[j]) ** 2)
        datab.append(res)
        res = []
    res = []
    for i in patientdatas:
        res.append(i[0])
        for j in range(1, 90):
            res.append((middledata[j] - i[j]) ** 2)
        datab.append(res)
        res = []

    datanolbir = []
    res = []

    for i in datab:
        res.append(i[0])
        for j in range(1, 90):
            if adata[j] != 0:
                if i[j] != 0:
                    nolbir = i[j] / adata[j]
                    if nolbir <= 1:
                        res.append(1)
                    else:
                        res.append(0)
                else:
                    res.append(0)
            else:
                res.append(0)
        datanolbir.append(res)
        res = []
    for i in datanolbir:
        nolbir, bol = Nolbir.objects.get_or_create(x1=i[1],
                                                   x2=i[2],
                                                   x3=i[3],
                                                   x4=i[4],
                                                   x5=i[5],
                                                   x6=i[6],
                                                   x7=i[7],
                                                   x8=i[8],
                                                   x9=i[9],
                                                   x10=i[10],
                                                   x11=i[11],
                                                   x12=i[12],
                                                   x13=i[13],
                                                   x14=i[14],
                                                   x15=i[15],
                                                   x16=i[16],
                                                   x17=i[17],
                                                   x18=i[18],
                                                   x19=i[19],
                                                   x20=i[20],
                                                   x21=i[21],
                                                   x22=i[22],
                                                   x23=i[23],
                                                   x24=i[24],
                                                   x25=i[25],
                                                   x26=i[26],
                                                   x27=i[27],
                                                   x28=i[28],
                                                   x29=i[29],
                                                   x30=i[30],
                                                   x31=i[31],
                                                   x32=i[32],
                                                   x33=i[33],
                                                   x34=i[34],
                                                   x35=i[35],
                                                   x36=i[36],
                                                   x37=i[37],
                                                   x38=i[38],
                                                   x39=i[39],
                                                   x40=i[40],
                                                   x41=i[41],
                                                   x42=i[42],
                                                   x43=i[43],
                                                   x44=i[44],
                                                   x45=i[45],
                                                   x46=i[46],
                                                   x47=i[47],
                                                   x48=i[48],
                                                   x49=i[49],
                                                   x50=i[50],
                                                   x51=i[51],
                                                   x52=i[52],
                                                   x53=i[53],
                                                   x54=i[54],
                                                   x55=i[55],
                                                   x56=i[56],
                                                   x57=i[57],
                                                   x58=i[58],
                                                   x59=i[59],
                                                   x60=i[60],
                                                   x61=i[61],
                                                   x62=i[62],
                                                   x63=i[63],
                                                   x64=i[64],
                                                   x65=i[65],
                                                   x66=i[66],
                                                   x67=i[67],
                                                   x68=i[68],
                                                   x69=i[69],
                                                   x70=i[70],
                                                   x71=i[71],
                                                   x72=i[72],
                                                   x73=i[73],
                                                   x74=i[74],
                                                   x75=i[75],
                                                   x76=i[76],
                                                   x77=i[77],
                                                   x78=i[78],
                                                   x79=i[79],
                                                   x80=i[80],
                                                   x81=i[81],
                                                   x82=i[82],
                                                   x83=i[83],
                                                   x84=i[84],
                                                   x85=i[85],
                                                   x86=i[86],
                                                   x87=i[87],
                                                   x88=i[88],
                                                   x89=i[89], )
        if type(i[0]) is int:
            etalon = Etalon.objects.get(id=i[0])
            etalon.nolbir_id = nolbir.id
            etalon.save()
        else:
            id = i[0].split(' ')
            id = id[-1].split('_')
            id = id[1]
            patient = Patient.objects.get(id=int(id))
            patient.nolbir = nolbir
            patient.save()

    return {'data': dataetalon,
            'adata': adata,
            'patient': datapatient,
            'middledata': middledata,
            'datab': datab,
            'datanolbir': datanolbir
            }


def normallashtirish(clas, patient_id):
    etalons = Etalon.objects.filter(clas=clas)
    patients = Patient.objects.filter(id=patient_id)
    etalondatas = []
    res = []
    sanoq = 1
    for i in etalons:
        res.append(f"{sanoq} {i.id}")
        res.append(i.x1)
        res.append(i.x2)
        res.append(i.x3)
        res.append(i.x4)
        res.append(i.x5)
        res.append(i.x6)
        res.append(i.x7)
        res.append(i.x8)
        res.append(i.x9)
        res.append(i.x10)
        res.append(i.x11)
        res.append(i.x12)
        res.append(i.x13)
        res.append(i.x14)
        res.append(i.x15)
        res.append(i.x16)
        res.append(i.x17)
        res.append(i.x18)
        res.append(i.x19)
        res.append(i.x20)
        res.append(i.x21)
        res.append(i.x22)
        res.append(i.x23)
        res.append(i.x24)
        res.append(i.x25)
        res.append(i.x26)
        res.append(i.x27)
        res.append(i.x28)
        res.append(i.x29)
        res.append(i.x30)
        res.append(i.x31)
        res.append(i.x32)
        res.append(i.x33)
        res.append(i.x34)
        res.append(i.x35)
        res.append(i.x36)
        res.append(i.x37)
        res.append(i.x38)
        res.append(i.x39)
        res.append(i.x40)
        res.append(i.x41)
        res.append(i.x42)
        res.append(i.x43)
        res.append(i.x44)
        res.append(i.x45)
        res.append(i.x46)
        res.append(i.x47)
        res.append(i.x48)
        res.append(i.x49)
        res.append(i.x50)
        res.append(i.x51)
        res.append(i.x52)
        res.append(i.x53)
        res.append(i.x54)
        res.append(i.x55)
        res.append(i.x56)
        res.append(i.x57)
        res.append(i.x58)
        res.append(i.x59)
        res.append(i.x60)
        res.append(i.x61)
        res.append(i.x62)
        res.append(i.x63)
        res.append(i.x64)
        res.append(i.x65)
        res.append(i.x66)
        res.append(i.x67)
        res.append(i.x68)
        res.append(i.x69)
        res.append(i.x70)
        res.append(i.x71)
        res.append(i.x72)
        res.append(i.x73)
        res.append(i.x74)
        res.append(i.x75)
        res.append(i.x76)
        res.append(i.x77)
        res.append(i.x78)
        res.append(i.x79)
        res.append(i.x80)
        res.append(i.x81)
        res.append(i.x82)
        res.append(i.x83)
        res.append(i.x84)
        res.append(i.x85)
        res.append(i.x86)
        res.append(i.x87)
        res.append(i.x88)
        res.append(i.x89)
        etalondatas.append(res)
        res = []
        sanoq += 1

    for i in patients:
        res = []
        res.append(f"{i.last_name}  {i.first_name} {i.id}")
        res.append(i.x1)
        res.append(i.x2)
        res.append(i.x3)
        res.append(i.x4)
        res.append(i.x5)
        res.append(i.x6)
        res.append(i.x7)
        res.append(i.x8)
        res.append(i.x9)
        res.append(i.x10)
        res.append(i.x11)
        res.append(i.x12)
        res.append(i.x13)
        res.append(i.x14)
        res.append(i.x15)
        res.append(i.x16)
        res.append(i.x17)
        res.append(i.x18)
        res.append(i.x19)
        res.append(i.x20)
        res.append(i.x21)
        res.append(i.x22)
        res.append(i.x23)
        res.append(i.x24)
        res.append(i.x25)
        res.append(i.x26)
        res.append(i.x27)
        res.append(i.x28)
        res.append(i.x29)
        res.append(i.x30)
        res.append(i.x31)
        res.append(i.x32)
        res.append(i.x33)
        res.append(i.x34)
        res.append(i.x35)
        res.append(i.x36)
        res.append(i.x37)
        res.append(i.x38)
        res.append(i.x39)
        res.append(i.x40)
        res.append(i.x41)
        res.append(i.x42)
        res.append(i.x43)
        res.append(i.x44)
        res.append(i.x45)
        res.append(i.x46)
        res.append(i.x47)
        res.append(i.x48)
        res.append(i.x49)
        res.append(i.x50)
        res.append(i.x51)
        res.append(i.x52)
        res.append(i.x53)
        res.append(i.x54)
        res.append(i.x55)
        res.append(i.x56)
        res.append(i.x57)
        res.append(i.x58)
        res.append(i.x59)
        res.append(i.x60)
        res.append(i.x61)
        res.append(i.x62)
        res.append(i.x63)
        res.append(i.x64)
        res.append(i.x65)
        res.append(i.x66)
        res.append(i.x67)
        res.append(i.x68)
        res.append(i.x69)
        res.append(i.x70)
        res.append(i.x71)
        res.append(i.x72)
        res.append(i.x73)
        res.append(i.x74)
        res.append(i.x75)
        res.append(i.x76)
        res.append(i.x77)
        res.append(i.x78)
        res.append(i.x79)
        res.append(i.x80)
        res.append(i.x81)
        res.append(i.x82)
        res.append(i.x83)
        res.append(i.x84)
        res.append(i.x85)
        res.append(i.x86)
        res.append(i.x87)
        res.append(i.x88)
        res.append(i.x89)
        etalondatas.append(res)

    maxdata = []
    maxdata.append('Max')
    for j in range(1, 90):
        max = 0
        for i in etalondatas:
            if i[j] > max:
                max = i[j]
        maxdata.append(max)
    normaldatas = []
    res = []
    sanoq = 1
    for i in etalondatas:
        res.append(i[0])
        id = i[0].split(' ')[-1]
        for j in range(1, 90):
            if maxdata[j] != 0:
                if maxdata[j] != 0:
                    res.append(i[j] / maxdata[j])
                else:
                    res.append(0)
            else:
                res.append(0)
        normallash, bol = Normallash.objects.get_or_create(x1=res[1],
                                                           x2=res[2],
                                                           x3=res[3],
                                                           x4=res[4],
                                                           x5=res[5],
                                                           x6=res[6],
                                                           x7=res[7],
                                                           x8=res[8],
                                                           x9=res[9],
                                                           x10=res[10],
                                                           x11=res[11],
                                                           x12=res[12],
                                                           x13=res[13],
                                                           x14=res[14],
                                                           x15=res[15],
                                                           x16=res[16],
                                                           x17=res[17],
                                                           x18=res[18],
                                                           x19=res[19],
                                                           x20=res[20],
                                                           x21=res[21],
                                                           x22=res[22],
                                                           x23=res[23],
                                                           x24=res[24],
                                                           x25=res[25],
                                                           x26=res[26],
                                                           x27=res[27],
                                                           x28=res[28],
                                                           x29=res[29],
                                                           x30=res[30],
                                                           x31=res[31],
                                                           x32=res[32],
                                                           x33=res[33],
                                                           x34=res[34],
                                                           x35=res[35],
                                                           x36=res[36],
                                                           x37=res[37],
                                                           x38=res[38],
                                                           x39=res[39],
                                                           x40=res[40],
                                                           x41=res[41],
                                                           x42=res[42],
                                                           x43=res[43],
                                                           x44=res[44],
                                                           x45=res[45],
                                                           x46=res[46],
                                                           x47=res[47],
                                                           x48=res[48],
                                                           x49=res[49],
                                                           x50=res[50],
                                                           x51=res[51],
                                                           x52=res[52],
                                                           x53=res[53],
                                                           x54=res[54],
                                                           x55=res[55],
                                                           x56=res[56],
                                                           x57=res[57],
                                                           x58=res[58],
                                                           x59=res[59],
                                                           x60=res[60],
                                                           x61=res[61],
                                                           x62=res[62],
                                                           x63=res[63],
                                                           x64=res[64],
                                                           x65=res[65],
                                                           x66=res[66],
                                                           x67=res[67],
                                                           x68=res[68],
                                                           x69=res[69],
                                                           x70=res[70],
                                                           x71=res[71],
                                                           x72=res[72],
                                                           x73=res[73],
                                                           x74=res[74],
                                                           x75=res[75],
                                                           x76=res[76],
                                                           x77=res[77],
                                                           x78=res[78],
                                                           x79=res[79],
                                                           x80=res[80],
                                                           x81=res[81],
                                                           x82=res[82],
                                                           x83=res[83],
                                                           x84=res[84],
                                                           x85=res[85],
                                                           x86=res[86],
                                                           x87=res[87],
                                                           x88=res[88],
                                                           x89=res[89])
        if sanoq != len(etalondatas):
            etalon = Etalon.objects.get(id=id)
            etalon.normallash = normallash
            etalon.save()
        else:
            patient = patients.last()
            patient.normallash = normallash
            patient.save()
        normaldatas.append(res)
        res = []
        sanoq += 1

    return {'normaldatas': normaldatas,
            'maxdata': maxdata,
            'data': etalondatas
            }

@admins_only
def nolbirView(request):
    if request.method == "POST":
        data = json.loads(request.body)
        clas = data['class']
        classnorm = Clas.objects.get(name = clas)
        patient = Patient.objects.filter(status='done').last()
        normallashtirish(classnorm, patient.id)
        data = nolbirlashtirish(clas, patient.id)
        nolbirdata = data['datanolbir']
        book = Workbook()
        sheet = book.active
        sheet.cell(1, 1).value = f"{patient.last_name}  {clas}"
        for i in range(len(nolbirdata)):
            for j in range(len(nolbirdata[i])):
                sheet.cell(row=i + 2, column=j + 1).value = nolbirdata[i][j]
        book.save(f'media/nolbir {clas} {patient.last_name}.xlsx')
        data['document'] = f'/media/nolbir {clas} {patient.last_name}.xlsx'
        return JsonResponse({'data': data})
    return render(request, 'nolbir.html', {})





@admins_only
def normView(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        clas = Clas.objects.get(name=data['class'])
        patients = Patient.objects.filter(status='done').last()
        patients = Patient.objects.filter(id=patients.id)
        normal_response = normallashtirish(clas, patients.last().id)
        etalondatas = normal_response['data']
        maxdata = normal_response['maxdata']
        normaldatas = normal_response['normaldatas']

        # writing to excel
        book = Workbook()
        sheet = book.active
        sheet.cell(1, 1).value = f"{patients.last().last_name}  {data['class']}"
        for i in range(len(normaldatas)):
            for j in range(len(normaldatas[i])):
                sheet.cell(row=i + 2, column=j + 1).value = normaldatas[i][j]
        book.save(f'media/normallashtirish {patients.last().last_name}.xlsx')

        return JsonResponse({'data': etalondatas,
                             'maxdata': maxdata,
                             'normaldata': normaldatas,
                             'document': f'/media/normallashtirish {patients.last().last_name}.xlsx'
                             })

    return render(request, 'normal.html', {})


def sinflashtirish(patient_id):
    data = {}
    for i in range(1, 6):
        clas = Clas.objects.get(name = f"{i}-class")
        print(f'{i}-class uchun sinflashtirish boshlandi')
        simptokomplexs = Simptokompleks.objects.all()
        datasend = []
        Lsimptokompleks.objects.all().delete()
        normallashtirish(clas, patient_id)
        nolbirlashtirish(f"{i}-class", patient_id)
        for simptokomplex in simptokomplexs:
            datazeroneid = []
            simpto = model_to_dict(simptokomplex)
            sanoql = 0
            for j in range(1, 90):
                if simpto[f'x{j}'] == 1:
                    datazeroneid.append(j)
                    sanoql += 1
            lsimpto, bol = Lsimptokompleks.objects.get_or_create(L=sanoql)
            etalons = Etalon.objects.filter(clas__name=f"{i}-class")
            datazerones = []
            for etalon in etalons:
                zerones = model_to_dict(etalon.nolbir)
                res = []

                for j in datazeroneid:
                    res.append(zerones[f'x{j}'])
                datazerones.append(res)
            patientzerone = []
            patient = Patient.objects.get(id=patient_id)
            patientdatas = model_to_dict(patient.nolbir)
            for j in datazeroneid:
                patientzerone.append(patientdatas[f'x{j}'])
            sanoq = 0
            for j in range(len(datazerones)):
                for k in range(len(patientzerone)):
                    if patientzerone[k] - datazerones[j][k] == 0:
                        sanoq += 1
            datasend.append([sanoq / len(datazerones), lsimpto.L])
        data[f"class{i}"] = datasend
        print(f'{i}-class uchun sinflashtirish yakunlandi')
    return data


@admins_only
def sinfview(request):
    if request.method == 'POST':
        patient_id = json.loads(request.body)
        patient_id = patient_id['patient_id']
        data = sinflashtirish(patient_id)
        return JsonResponse({'data': data})

    patient = Patient.objects.filter(status='done').last()
    context = {
        'patients': patient
    }

    return render(request, 'sinflashtirish.html', context)


@admins_only
def normresultView(request):
    if request.method == 'POST':
        patient = Patient.objects.filter(status='done').last()
        data = sinflashtirish(patient.id)
        lsimpto = Lsimptokompleks.objects.all()
        dataresponse = {}
        for i in range(1, 6):
            ldatas = data[f'class{i}']
            classdata = {}
            for j in lsimpto:
                classdata[f'L{j.L}'] = [0, 0]
            for j in ldatas:
                if j[1] != 0:
                    classdata[f'L{j[1]}'][0] += j[0] / j[1]
                classdata[f'L{j[1]}'][1] += 1

            for j in lsimpto:
                if classdata[f'L{j.L}'][1]:
                    classdata[f'L{j.L}'] = classdata[f'L{j.L}'][0] / classdata[f'L{j.L}'][1]
                else:
                    classdata[f'L{j.L}'] = 0
            dataresponse[f'class{i}'] = classdata
        print('jarayonda')
        Ldata = []
        for i in lsimpto:
            Ldata.append(i.L)
        maxium = []
        for i in range(1, 6):
            res = 0
            for j in Ldata:
                res += dataresponse[f'class{i}'][f'L{j}']
            maxium.append(res)
        maxvalue = max(maxium)
        maxindex = maxium.index(maxvalue)
        illname = Clas_ill_names.objects.get(clas__name=f'{maxindex + 1}-class').name
        patient.program_diagnos = illname
        patient.save()
        return JsonResponse({'data': dataresponse,
                             'Ldata': Ldata,
                             'max': maxium,
                             'ill': [illname, maxindex + 1]})

    return render(request, 'adminresultpage.html', {})


@doctors_only
def lastPatientView(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        patient = Patient.objects.get(id=data['patient_id'])
        patient.__dict__.update({f'{data["name"]}': float(data['value'])})
        patient.save()

        return JsonResponse({'data': 'ok'})
    patients = Patient.objects.filter(Q(status='begin') | Q(status='progress'), pat_doctor_id=request.user.id)

    return render(request, 'lastpatients.html', {'patients': patients})


@doctors_only
def addPatientView(request):
    if request.method == 'POST':
        last_name = request.POST.get('last_name')
        first_name = request.POST.get('first_name')
        middle_name = request.POST.get('middle_name')
        birthday = request.POST.get('dob')
        home_phone = request.POST.get('home_phone')
        work_phone = request.POST.get('work_phone')
        region = request.POST.get('viloyat')
        city = request.POST.get('tuman')
        adres = request.POST.get('address')
        gen = request.POST.get('gen')
        work_place = request.POST.get('work_place')
        patient = Patient.objects.create(last_name=last_name, first_name=first_name,
                                         middle_name=middle_name, birthday=birthday, home_phone=home_phone,
                                         work_phone=work_phone, region=region, city=city, adress=adres, gen=gen,
                                         work_place=work_place, status='begin')

        patient.pat_doctor_id = request.user.id
        patient.save()

        return redirect(f'/patient/{patient.id}')

    return render(request, 'add-patient.html')


@doctors_only
def diagnosedPatientView(request):
    patients = Patient.objects.filter(status='done', pat_doctor_id=request.user.id)
    data = []
    sanoq = 1

    for i in patients:

        changedata = model_to_dict(i)
        changedata['created_data'] = i.created_data.strftime("%Y-%m-%d")
        pdata = {}
        for key in changedata.keys():
            keydata = key[1:] + key[0]
            pdata[keydata] = changedata[key]
        document = Document('Kasallik_varaqasi_namuna.docx')
        for p in document.paragraphs:
            for key in changedata.keys():
                if key in p.text:
                    inline = p.runs
                    for item in inline:
                        if key == item.text:
                            item.text = item.text.replace(key, f'{changedata[key]}')
                    p.text.replace(f'{key}', f'{changedata[key]}')
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key in pdata.keys():
                            if key == paragraph.text:
                                paragraph.text = paragraph.text.replace(f"{key}", f"{int(pdata[key])}")
        document.save(f'media/word/{i.last_name} {i.first_name}.docx')
        data.append({
            'sanoq': sanoq,
            'name': f'{i.last_name}  {i.first_name}',
            'birthday': f'{i.birthday}',
            'created_data': i.created_data,
            'final_diagnos': i.program_diagnos,
            'document': f'/media/word/{i.last_name} {i.first_name}.docx'
        })
        sanoq += 1
    return render(request, 'diagnosed-patient.html', {'patients': data})


@doctors_only
def PatientView(request, id):
    if request.method == 'POST':
        calculus_status = Patient_informs.objects.all().first()
        data = json.loads(request.body)
        patient_id = data['patient_id']
        patient = Patient.objects.get(id=patient_id)
        patientcheck = model_to_dict(patient)
        bol = True
        for i in range(1, 90):
            if patientcheck[f'x{i}'] == None:
                bol = False
        if bol:
            patient.status = 'progress'
            calculus_status.calculus_status = False
            calculus_status.save()
            patient.save()
            for i in range(1, 6):
                classnorm = Clas.objects.get(name=f'{i}-class')
                normallashtirish(classnorm, patient.id)
                nolbirlashtirish(f'{i}-class', patient_id)
            data = sinflashtirish(patient.id)
            lsimpto = Lsimptokompleks.objects.all()
            dataresponse = {}
            for i in range(1, 6):
                ldatas = data[f'class{i}']
                classdata = {}
                for j in lsimpto:
                    classdata[f'L{j.L}'] = [0, 0]
                for j in ldatas:
                    if j[1] != 0:
                        classdata[f'L{j[1]}'][0] += j[0] / j[1]
                    classdata[f'L{j[1]}'][1] += 1
                for j in lsimpto:
                    if classdata[f'L{j.L}'][1]:
                        classdata[f'L{j.L}'] = classdata[f'L{j.L}'][0] / classdata[f'L{j.L}'][1]
                    else:
                        classdata[f'L{j.L}'] = 0
                dataresponse[f'class{i}'] = classdata

            Ldata = []
            for i in lsimpto:
                Ldata.append(i.L)
            maxium = []
            for i in range(1, 6):
                res = 0
                for j in Ldata:
                    res += dataresponse[f'class{i}'][f'L{j}']
                maxium.append(res)
            maxvalue = max(maxium)
            maxindex = maxium.index(maxvalue)
            illname = Clas_ill_names.objects.get(clas__name=f'{maxindex + 1}-class').name
            doctor = patient.doctor
            tashxis = patient.doctor_diagnos
            partboss = patient.partboss
            maindoctor = patient.maindoctor
            patient.program_diagnos = illname
            if doctor == None:
                doctor = ''
            if tashxis == None:
                tashxis = ''
            if partboss == None:
                partboss = ''
            if maindoctor == None:
                maindoctor = ''
            patient.save()
            calculus_status.calculus_status = True
            calculus_status.save()
            return JsonResponse({'ill': [illname, maxindex + 1],
                                 'doctor': doctor,
                                 'tashxis': tashxis,
                                 'partboss': partboss,
                                 'maindoctor': maindoctor
                                 })
        else:
            return JsonResponse({'data': 'Hali barcha qismlar to\'ldirilmagan'})
    try:
        patient = Patient.objects.get(id=id)
        if patient.status == 'done':
            return redirect('lastpatient')
    except:
        return HttpResponse('404 page not found')
    return render(request, 'Patient.html', {'patient': patient})


@doctors_only
def patient_render_shikoyat(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        id = data['id']
        patient_id = data['patient_id']
        data = model_to_dict(Patient_informs.objects.get(id=1))
        patient = model_to_dict(Patient.objects.get(id=patient_id))
        datasend = []
        res = []
        comment = ''
        birlik = ''
        if id == 1:
            datasend.append('Беморнинг шикояти ва анамнез маълумотлари киритиш ойнаси')
            comment = 'Беморда мазкур касаллик белгиси(симптоми) кузатилган бўлса 1, акс ҳолда 0 рақами танланади.'
            for i in range(1, 20):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        elif id == 2:
            birlik = ['минут', 'минут', 'мм(см устун)', 'мм(см устун)', '%']
            datasend.append('Беморнинг бирламчи кўрик натижаларини киритиш ойнаси')
            for i in range(20, 25):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                res.append(i - 20)
                datasend.append(res)
                res = []
        elif id == 3:
            birlik = ['г/л', 'г/л', '-', 'г/л', '%', 'мм/соат', 'мкл', '%', '%', '%', 'минут', 'минут', 'сек', 'даража',
                      'г/л', '%', 'мин', '%', 'МЕ/л', 'МЕ/л', 'мкмоль/л', 'мкмоль/л', 'мкмоль/л', 'мкмоль/л', 'Ед/л',
                      'ммоль/л']
            datasend.append("Беморнинг лаборатория натижаларини киритиш ойнаси")
            for i in range(25, 51):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        elif id == 4:
            datasend.append("Беморнинг ЭКГ натижаларини киритиш ойнаси")
            comment = 'Беморда мазкур касаллик белгиси(симптоми) кузатилган бўлса 1, акс ҳолда 0 рақами танланади.'
            for i in range(51, 64):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        elif id == 5:
            datasend.append("Беморнинг ЭХО кардиография  натижаларини киритиш ойнаси")
            birlik = ['мм', 'мм', 'мм', 'мм', 'мм', 'мл', 'мл', 'мл', 'мл/сек']
            for i in range(64, 73):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        elif id == 6:
            datasend.append("Беморда мавжуд ёндош касалликларини киритиш ойнаси")
            comment = 'Беморда мазкур касаллик белгиси(симптоми) кузатилган бўлса 1, акс ҳолда 0 рақами танланади.'
            for i in range(73, 80):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        elif id == 7:
            datasend.append("Беморнинг текширув натижалари таҳлили ва якуний ташхислаш ойнаси")
            comment = 'Беморда мазкур касаллик белгиси(симптоми) кузатилган бўлса 1, акс ҳолда 0 рақами танланади.'
            for i in range(80, 90):
                res.append(f'x{i}')
                res.append(data[f'x{i}'])
                res.append(patient[f'x{i}'])
                datasend.append(res)
                res = []
        return JsonResponse({'data': datasend,
                             'comment': comment,
                             'birlik': birlik})


@doctors_only
def final_result(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        patient = Patient.objects.get(id=data['patient_id'])
        patient.__dict__.update({f'{data["name"]}': data['value']})
        patient.save()
        return JsonResponse({'data': 'ok'})


@doctors_only
def save_patient(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        patient = Patient.objects.get(id=data['patient_id'])
        if patient.doctor and patient.maindoctor and patient.doctor_diagnos and patient.partboss:
            patient.status = 'done'
            patient.save()
            return JsonResponse({'data': 'Bemor muaffaqiyantli saqlandi'})
        else:
            return JsonResponse({
                'error': "Hali belgilangan barcha qismlar to'ldirilgani yo'q iltimos malumotlarni to'ldirib qaya harakat qilib ko'ring"})


@login_required(redirect_field_name='login', login_url='/login/')
def user_page(request):
    if request.method == 'POST':
        passlast = request.POST.get('pass1')
        passnew1 = request.POST.get('newpass1')
        passnew2 = request.POST.get('newpass2')
        if passnew1 != passnew2:
            return render(request, 'userpage.html',
                          {'error': "kiritgan passwordlaringiz bir xil emas iltimos qayta tekshirib ko'ring"})
        else:
            user = authenticate(request, username=request.user.username, password=passlast)
            if user:
                user.set_password(passnew2)
                user.save()
                return render(request, 'userpage.html', {'error': "Sizning parolingiz muaffaqiyatli o'zgartirildi"})
            else:
                return render(request, 'userpage.html', {
                    'error': 'Sizning Joriy parolingiz xato kiritilgan iltimos qayta harakat qilib ko\'ring'})
    return render(request, 'userpage.html', {'error': ''})
